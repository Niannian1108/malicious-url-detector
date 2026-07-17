from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TECH_OUT = Path(r"C:/Users/Chong You Xin/Documents/FYP 2 Documents/Resubmission Edit/1211108623_CHONG YOU XIN_TECHNICALPAPER_ENHANCED_HIGHLIGHTED.docx")


def insert_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = "Paragraph"
    except Exception:
        pass
    run = new_para.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_para


def main() -> None:
    doc = Document(str(TECH_OUT))
    citations = {
        "TABLE 1. Summary of the implemented dataset.": "Table 1 summarises the corrected model-development dataset and the held-out hard-negative evaluation set used in this study.",
        "TABLE 2. Stratified train-test distribution.": "Table 2 presents the stratified train-test distribution used for the main machine-learning evaluation.",
        "TABLE 3. Engineered feature categories.": "Table 3 lists the engineered feature groups extracted from each URL before model training and inference.",
        "TABLE 4. Machine learning model comparison on the held-out test set.": "Table 4 compares the four supervised learning models using the corrected held-out test set.",
        "TABLE 5. Runtime risk classification and response.": "Table 5 shows how backend model outputs are translated into low, medium and high-risk browser responses.",
        "TABLE 6. Selected feature importance values.": "Table 6 reports selected Gradient Boosting feature-importance values used to interpret the final classifier.",
        "TABLE 7. Deployment threshold and hard-negative evaluation results.": "Table 7 summarises deployment-threshold performance and held-out hard-negative false-positive evaluation.",
    }
    existing = "\n".join(p.text for p in doc.paragraphs)
    for caption, citation in citations.items():
        if citation in existing:
            continue
        for p in doc.paragraphs:
            if p.text.strip() == caption:
                insert_before(p, citation)
                break
    doc.save(str(TECH_OUT))
    print(f"Patched table citations in {TECH_OUT}")


if __name__ == "__main__":
    main()
