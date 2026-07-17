from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:/Users/Chong You Xin")
RESUB = ROOT / "Documents/FYP 2 Documents/Resubmission Edit"
TECH_SRC = ROOT / "Documents/FYP 2 Documents/Technical Report/1211108623_CHONG YOU XIN_TECHNICALPAPER.docx"
JIWE_TEMPLATE = RESUB / "JIWE_Template-new_jan2026-OTH.docx"

TECH_OUT = RESUB / "1211108623_CHONG YOU XIN_TECHNICALPAPER_ENHANCED_HIGHLIGHTED.docx"
JIWE_OUT = RESUB / "JIWE_Malicious_Website_Detection_CHONG_YOU_XIN_FILLED.docx"

LATEST = ROOT / "Documents/FYP 2 Documents/Latest Screenshot"

SCREENSHOTS = {
    "extension": LATEST / "Extension Card Loaded in Google Chrome.png",
    "fastapi": LATEST / "FastAPI Page.png",
    "block": LATEST / "Block Page Warning.png",
    "medium": LATEST / "Medium Risk Warning Popup.png",
    "proceed": LATEST / "Proceed Page.png",
}


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_highlighted(paragraph: Paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_paragraph_text(paragraph: Paragraph, text: str, highlight: bool = False) -> None:
    style = paragraph.style
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.style = style
    run = paragraph.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def add_highlighted_citation_before_caption(doc: Document, caption: str, citation: str) -> None:
    cap = find_paragraph(doc, caption)
    p = paragraph_after(doc.paragraphs[doc.paragraphs.index(cap) - 1], "", "Paragraph")
    # paragraph_after inserts after previous paragraph, therefore just before caption.
    add_highlighted(p, citation)


def add_picture_after(paragraph: Paragraph, image_path: Path, caption: str, width_inches: float = 5.7) -> Paragraph:
    img_p = paragraph_after(paragraph, "", "Paragraph")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap_p = paragraph_after(img_p, "", "Table Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_highlighted(cap_p, caption)
    return cap_p


def enhance_technical_paper() -> None:
    shutil.copyfile(TECH_SRC, TECH_OUT)
    doc = Document(str(TECH_OUT))
    doc.core_properties.title = "Machine Learning Based Malicious Website Detection and Response System"
    doc.core_properties.author = "Chong You Xin"
    doc.core_properties.last_modified_by = "Chong You Xin"

    table_citations = {
        "TABLE 1. Summary of the implemented dataset.": "Table 1 summarises the corrected model-development dataset and the held-out hard-negative evaluation set used in this study.",
        "TABLE 2. Stratified train-test distribution.": "Table 2 presents the stratified train-test distribution used for the main machine-learning evaluation.",
        "TABLE 3. Engineered feature categories.": "Table 3 lists the engineered feature groups extracted from each URL before model training and inference.",
        "TABLE 4. Machine learning model comparison on the held-out test set.": "Table 4 compares the four supervised learning models using the corrected held-out test set.",
        "TABLE 5. Runtime risk classification and response.": "Table 5 shows how backend model outputs are translated into low, medium and high-risk browser responses.",
        "TABLE 6. Selected feature importance values.": "Table 6 reports selected Gradient Boosting feature-importance values used to interpret the final classifier.",
        "TABLE 7. Deployment threshold and hard-negative evaluation results.": "Table 7 summarises deployment-threshold performance and held-out hard-negative false-positive evaluation.",
    }
    for caption, citation in table_citations.items():
        try:
            add_highlighted_citation_before_caption(doc, caption, citation)
        except ValueError:
            pass

    system_design = find_paragraph(doc, "SYSTEM DESIGN")
    p = paragraph_after(system_design, "", "Paragraph")
    add_highlighted(
        p,
        "The implemented prototype includes a Chrome extension interface and a local FastAPI backend. "
        "Figure 1 shows the loaded browser extension card used to access the detector during testing, while Figure 2 shows the FastAPI documentation page used to verify backend availability and prediction endpoints.",
    )
    cap = add_picture_after(p, SCREENSHOTS["extension"], "Figure 1. Chrome extension loaded in the browser.")
    cap = add_picture_after(cap, SCREENSHOTS["fastapi"], "Figure 2. FastAPI backend documentation and endpoint interface.")

    response_anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("The response layer specifies three severities"):
            response_anchor = p
            break
    if response_anchor is not None:
        p = paragraph_after(response_anchor, "", "Paragraph")
        add_highlighted(
            p,
            "The browser response layer was tested through high-risk blocking, medium-risk warning and user override scenarios. "
            "Figure 3 shows the high-risk warning page, Figure 4 shows the medium-risk caution popup, and Figure 5 shows the proceed page used when a user chooses to continue after a warning.",
        )
        cap = add_picture_after(p, SCREENSHOTS["block"], "Figure 3. High-risk warning page displayed for a blocked malicious URL.")
        cap = add_picture_after(cap, SCREENSHOTS["medium"], "Figure 4. Medium-risk caution popup shown for suspicious but non-blocked URLs.")
        add_picture_after(cap, SCREENSHOTS["proceed"], "Figure 5. Proceed page shown after a one-time user override.")

    results = find_paragraph(doc, "RESULTS AND EVALUATION")
    p = paragraph_after(results, "", "Paragraph")
    add_highlighted(
        p,
        "The additional screenshots demonstrate that the proposed system is not limited to an offline classifier. "
        "The trained model, backend inference service, browser extension, warning page and user override flow work together as a complete detection-and-response prototype.",
    )

    doc.save(str(TECH_OUT))


def clear_body(doc: Document) -> None:
    for tbl in list(doc.tables):
        remove_element(tbl._element)
    for p in list(doc.paragraphs):
        remove_element(p._element)


def set_run_font(run, size: int = 10, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc: Document, text: str = "", size: int = 10, bold: bool = False, italic: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_like(doc: Document, text: str) -> None:
    add_para(doc, text, size=10, bold=True)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_para(doc, caption, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        run = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(run, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(val)
            set_run_font(run)


def add_jiwe_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 5.6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_para(doc, caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def fill_jiwe_template() -> None:
    shutil.copyfile(JIWE_TEMPLATE, JIWE_OUT)
    doc = Document(str(JIWE_OUT))
    clear_body(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_para(doc, "Journal of Informatics and", size=12, bold=True)
    add_para(doc, "Web Engineering", size=12, bold=True)
    add_para(doc, "Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X", size=9)
    add_para(doc, "Machine Learning Based Malicious Website Detection and Response System", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Chong You Xin", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    abstract = (
        "Abstract - Malicious and phishing websites remain common web-based threats that exploit user trust, brand familiarity and rapidly changing URL structures. "
        "This paper presents a machine learning based malicious website detection and response system that combines engineered URL features, trust-aware indicators, brand-mismatch detection, lightweight runtime page evidence, a FastAPI inference backend and a Chrome extension response layer. "
        "A corrected model-development dataset of 1,178 URLs was used, consisting of 878 benign URLs and 300 malicious URLs, while a separate 21-URL held-out hard-negative benign set was excluded from training and used only for false-positive evaluation. "
        "Twenty-two lexical, security, trust and brand-impersonation features were extracted from each URL. Four supervised models were compared, namely Logistic Regression, Support Vector Machine with RBF kernel, Random Forest and Gradient Boosting. "
        "Gradient Boosting achieved 98.31% accuracy, 98.28% malicious precision, 95.00% malicious recall, 96.61% F1-score and 99.24% ROC-AUC on the corrected main hold-out test split. "
        "At a 0.90 deployment threshold, the prototype achieved 93.33% malicious recall, 100.00% malicious precision, 0.00% main benign false positives and 0.00% held-out hard-negative false positives. "
        "The implemented browser prototype converts model decisions into low, medium and high-risk responses, including caution warnings and high-risk blocking pages."
    )
    add_para(doc, abstract, size=9)
    add_para(doc, "Keywords—Machine Learning, Malicious URL Detection, Phishing Detection, Chrome Extension, FastAPI, Cybersecurity", size=9, italic=True)

    add_heading_like(doc, "INTRODUCTION")
    add_para(doc, "Web-based attacks such as phishing and malicious websites remain effective because attackers can quickly create domains, imitate trusted brands and present familiar login or verification pages. Static blacklists and signature-based tools are useful, but they may react too slowly to newly created URLs. This study therefore proposes a machine learning based detection-and-response prototype that checks URL structure, domain-related features, trust-aware signals, brand mismatch and lightweight runtime page evidence.")
    add_para(doc, "The contribution of this work is not only the classifier, but also the integration of the classifier into a usable browser workflow. The Chrome extension sends URLs to a FastAPI backend, receives the risk result and presents an appropriate response to the user. Low-risk pages are allowed, medium-risk pages produce caution warnings, and high-risk pages are redirected to a warning page.")

    add_heading_like(doc, "LITERATURE REVIEW")
    add_para(doc, "Existing studies have explored lexical URL features, domain features, webpage content, external reputation signals and deep learning sequence models for malicious URL detection [1]-[5]. Classical models can remain competitive when feature engineering is carefully designed [6], [7]. Recent work also highlights adversarial issues such as evasion and label manipulation [8], [9]. However, many studies focus on offline classification accuracy, while fewer demonstrate how model results are translated into practical user-facing browser responses.")

    add_heading_like(doc, "RESEARCH METHODOLOGY")
    add_para(doc, "The dataset used in this study is summarised in Table 1. The held-out hard-negative set was excluded from model training so that false-positive evaluation is performed on unseen legitimate but suspicious-looking URLs.")
    add_table(doc, "Table 1. Dataset composition.", ["Class", "Number of URLs"], [["Benign used for model development", "878"], ["Malicious used for model development", "300"], ["Held-out hard-negative benign evaluation URLs", "21"]])
    add_para(doc, "The stratified train-test split for the main model-development data is shown in Table 2.")
    add_table(doc, "Table 2. Stratified train-test distribution.", ["Dataset", "Benign", "Malicious", "Total"], [["Training Set", "702", "240", "942"], ["Testing Set", "176", "60", "236"], ["Total model-development data", "878", "300", "1,178"]])
    add_para(doc, "Table 3 lists the four feature categories extracted from each URL before model training and inference.")
    add_table(doc, "Table 3. Engineered feature categories.", ["Category", "Features"], [["Lexical", "URL length, host length, path length, query length, dots, digits, hyphens, special characters, entropy, subdomains and path depth"], ["Security indicators", "HTTPS use, IP address use, punycode, executable path, suspicious TLD and suspicious keywords"], ["Trust-based", "Known trusted-domain recognition"], ["Brand impersonation", "Brand keyword and brand-domain mismatch indicators"]])

    add_heading_like(doc, "SYSTEM DESIGN AND PROTOTYPE IMPLEMENTATION")
    add_para(doc, "The implemented system separates browser interaction from backend inference. The Chrome extension monitors navigation and collects lightweight DOM signals, while the FastAPI backend extracts features, loads the deployed Gradient Boosting model, determines risk severity and logs prediction events. Figure 1 shows the loaded Chrome extension interface, and Figure 2 shows the FastAPI backend interface.")
    add_jiwe_figure(doc, SCREENSHOTS["extension"], "Figure 1. Chrome extension loaded in the browser.")
    add_jiwe_figure(doc, SCREENSHOTS["fastapi"], "Figure 2. FastAPI backend documentation and endpoint interface.")
    add_para(doc, "The response mechanism uses three risk levels as shown in Table 4. This allows the system to avoid unnecessary blocking when evidence is weak while still responding strongly to high-confidence malicious results.")
    add_table(doc, "Table 4. Runtime risk classification and response.", ["Risk Level", "Condition", "Response"], [["Low", "Benign prediction", "Allow navigation"], ["Medium", "Malicious prediction with confidence >= 0.70", "Display caution notification"], ["High", "Malicious prediction with confidence >= 0.90 and supporting evidence", "Display warning/block page"]])
    add_para(doc, "Figure 3 shows the high-risk warning page, Figure 4 shows the medium-risk caution popup, and Figure 5 shows the proceed page displayed after a one-time user override.")
    add_jiwe_figure(doc, SCREENSHOTS["block"], "Figure 3. High-risk warning page displayed for a blocked malicious URL.")
    add_jiwe_figure(doc, SCREENSHOTS["medium"], "Figure 4. Medium-risk caution popup shown for suspicious but non-blocked URLs.")
    add_jiwe_figure(doc, SCREENSHOTS["proceed"], "Figure 5. Proceed page shown after a one-time user override.")

    add_heading_like(doc, "RESULTS AND DISCUSSIONS")
    add_para(doc, "Table 5 compares the four evaluated machine learning models on the corrected held-out test set. Gradient Boosting and Random Forest achieved similar main accuracy, but Gradient Boosting was selected because it obtained the stronger ROC-AUC and deployment-threshold recall.")
    add_table(doc, "Table 5. Machine learning model comparison.", ["Model", "Accuracy", "Precision", "Recall", "F1-score", "ROC-AUC"], [["Gradient Boosting", "98.31%", "98.28%", "95.00%", "96.61%", "99.24%"], ["Random Forest", "98.31%", "98.28%", "95.00%", "96.61%", "98.71%"], ["SVM (RBF)", "97.03%", "93.44%", "95.00%", "94.21%", "97.18%"], ["Logistic Regression", "96.19%", "91.80%", "93.33%", "92.56%", "98.71%"]])
    add_para(doc, "Table 6 summarises the deployed threshold and held-out hard-negative evaluation results. The 0.90 threshold prioritises high-confidence blocking to reduce unnecessary disruption to legitimate browsing.")
    add_table(doc, "Table 6. Deployment threshold and hard-negative evaluation results.", ["Metric", "Result"], [["Main benign false-positive rate", "0.00%"], ["Main malicious recall", "93.33%"], ["Main malicious precision", "100.00%"], ["Held-out hard-negative URLs tested", "21"], ["Held-out hard-negative false-positive rate", "0.00%"], ["Software verification tests", "25 / 25 tests passed"]])
    add_para(doc, "The results show that a classical ensemble model with carefully engineered URL and trust-aware features can support a real-time malicious website detection prototype. The held-out hard-negative result is important because login and account-security pages may look suspicious lexically but are legitimate when hosted on trusted domains.")

    add_heading_like(doc, "CONCLUSION")
    add_para(doc, "This paper presented a machine learning based malicious website detection and response system that integrates URL classification with browser-level user protection. The corrected evaluation achieved strong main test-set performance and zero false positives on the held-out hard-negative benign set. Future work should expand the dataset, strengthen privacy controls for logging and reputation checks, and add richer evidence such as redirect-chain, certificate and page-content features.")

    add_heading_like(doc, "ACKNOWLEDGEMENT")
    add_para(doc, "The author thanks the Faculty of Information Science and Technology, Multimedia University, and Prof. Ts. Dr. Heng Swee Huay for guidance and support throughout the final-year project.")
    add_heading_like(doc, "FUNDING STATEMENT")
    add_para(doc, "The author received no funding from any party for the research and publication of this article.")
    add_heading_like(doc, "AUTHOR CONTRIBUTIONS")
    add_para(doc, "Chong You Xin: Conceptualization, Data Curation, Methodology, Software, Validation, Writing - Original Draft Preparation.")
    add_heading_like(doc, "CONFLICT OF INTERESTS")
    add_para(doc, "No conflict of interests were disclosed.")
    add_heading_like(doc, "ETHICS STATEMENTS")
    add_para(doc, "This work used publicly available URL data and did not involve human or animal subjects.")
    add_heading_like(doc, "DATA AVAILABILITY")
    add_para(doc, "The data underlying this study are derived from publicly available benign and malicious URL sources described in the article and project repository.")
    add_heading_like(doc, "REFERENCES")
    refs = [
        "[1] Y. Tian, Y. Yu, J. Sun, and Y. Wang, \"From past to present: A survey of malicious URL detection techniques, datasets, and code repositories,\" arXiv:2504.16449, 2025.",
        "[2] F. Turk and M. Kilicaslan, \"Malicious URL detection with advanced machine learning and optimization-supported deep learning models,\" Applied Sciences, vol. 15, p. 10090, 2025.",
        "[3] Q. E. Haq, M. H. Faheem, and I. Ahmad, \"Detecting phishing URLs based on a deep learning approach to prevent cyber-attacks,\" Applied Sciences, vol. 14, no. 22, p. 10086, 2024.",
        "[4] L. Zhang and Q. Yan, \"Detect malicious websites by building a neural network to capture global and local features of websites,\" Computers & Security, vol. 137, p. 103641, 2023.",
        "[5] T. Wu, M. Wang, Y. Xi, and Z. Zhao, \"Malicious URL detection model based on bidirectional gated recurrent unit and attention mechanism,\" Applied Sciences, vol. 12, p. 12367, 2022.",
        "[6] X. Guo, \"Evaluation on malicious URL detection with different features based on various machine learning algorithms,\" in Proceedings of the 12th International Conference on Data Science, pp. 543-550, SCITEPRESS, 2023.",
        "[7] M. Adam, S. F. Nasution, R. R. Simanungkalit, and I. H. Diansyah, \"Machine learning-driven detection of malicious URL: Comparative analysis of random forest and SVMs,\" JITE, vol. 8, 2024.",
        "[8] E. Nowroozi, A. M., M. Mohammadi, and M. Conti, \"An adversarial attack analysis on malicious advertisement URL detection framework,\" arXiv:2204.13172, 2022.",
        "[9] E. Nowroozi, N. Jadalla, S. Ghelichkhani, and A. Jolfaei, \"Mitigating label-flipping attacks in malicious URL detectors using ensemble trees,\" arXiv:2403.02995, 2024.",
        "[10] S. Aslam, H. Aslam, A. Manzoor, C. Hui, and A. Rasool, \"AntiPhishStack: LSTM-based stacked generalization model for optimized phishing URL detection,\" arXiv:2401.08947, 2024.",
    ]
    for ref in refs:
        add_para(doc, ref, size=9)

    doc.core_properties.title = "Machine Learning Based Malicious Website Detection and Response System"
    doc.core_properties.author = "Chong You Xin"
    doc.core_properties.last_modified_by = "Chong You Xin"
    doc.save(str(JIWE_OUT))


def main() -> None:
    for label, path in {"technical paper": TECH_SRC, "JIWE template": JIWE_TEMPLATE, **SCREENSHOTS}.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing {label}: {path}")
    enhance_technical_paper()
    fill_jiwe_template()
    print(f"Enhanced technical paper: {TECH_OUT}")
    print(f"Filled JIWE template: {JIWE_OUT}")


if __name__ == "__main__":
    main()
